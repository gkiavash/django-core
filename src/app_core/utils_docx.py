import pprint

import bs4
from copy import deepcopy
from collections import defaultdict

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_COLOR_INDEX

from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.lib import utils
from reportlab.lib import pagesizes

from app_process.rore import trp


FONT_NAME = 'Helvetica'
FONT_SIZE = 7
MAX_PAGE_WIDTH = 300


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
        # print(parent_elm.xml)
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def get_text_docx(file_path):
    """
    returns the text of docx file with considering table tags and cell orders
    """
    """
    The pdf is also created in order to have coordinates
    """
    document = Document(file_path)
    text = ''
    indexes = {}

    # PDF CREATION
    canvas_ = canvas.Canvas("hello.pdf", pagesize=pagesizes.A4)

    line_counter = 0
    lines_per_page = 40
    for block in iter_block_items(document):
        if isinstance(block, Paragraph):
            # txt_ += block.text + '\n'

            # PDF CREATION: needed to be split for line alignment in A4 paper
            lines_aligned = utils.simpleSplit(block.text, fontName=FONT_NAME, fontSize=FONT_SIZE, maxWidth=MAX_PAGE_WIDTH)
            for line_aligned in lines_aligned:
                for line in line_aligned.split('\n'):
                    textobject = canvas_.beginText(2 * cm, 29.7 * cm - (line_counter % lines_per_page * 0.5 * cm) - 2 * cm)

                    for word in line.split(' '):
                        txt_, indexes_ = docx_add_word(
                            word=word,
                            indexes_offset=len(text),
                            textobject=textobject
                        )
                        indexes.update(indexes_)
                        text += txt_
                        bbox = indexes_[list(indexes_.keys())[0]]['BoundingBox']
                        canvas_.rect(bbox.left, bbox.top + bbox.height, bbox.width, bbox.height, fill=False)

                    canvas_.drawText(textobject)

                    line_counter += 1
                    if line_counter % lines_per_page == lines_per_page-1:
                        canvas_.showPage()

        elif isinstance(block, Table):
            text += get_text_docx_tables_(
                table=block,
                line_counter=line_counter,
                lines_per_page=lines_per_page,
                canvas_=canvas_
            )

    canvas_.save()

    return text, indexes


def docx_add_word(word, textobject, indexes_offset=0,):
    # PDF: coordinates
    x_start = textobject.getCursor()[0]
    textobject.textOut(word + ' ')
    x_end = textobject.getCursor()[0]

    bbox = trp.BoundingBox(
        x_end - x_start,
        textobject._leading,
        x_start,
        textobject._y
    )  # width, height, left, top

    bboxes = {'BoundingBox': bbox, 'Text': word,}

    # TEXT
    txt_ = ''

    indexes_start = indexes_offset + len(txt_)
    txt_ += word + ' '
    indexes_end = indexes_start + len(word)
    indexes = {f'{indexes_start}_{indexes_end}': bboxes}

    return txt_, indexes


def get_text_docx_tables_(
        table,
        indexes_offset=0,
        line_counter=0,
        lines_per_page=40,
        canvas_=None,
        add_tags=True,
):
    """
    In docx library, merged cells are duplicated in iteration.
    Here, in order to avoid duplicated text, added cells' ids are stored in an array.
    If the cell is not in the array, then its text is added. Otherwise, the cell is skipped.
    """
    cells_added = []
    txt_ = '\n<table>\n' if add_tags else ''
    for row_index, row in enumerate(table.rows):
        txt_ += '<tr>' if add_tags else ''
        for cell_index, cell in enumerate(row.cells):
            if cell._tc not in cells_added:  # for merged cells
                if add_tags:
                    txt_ += f'<td>' if row_index == 0 else f'<td title="{table.rows[0].cells[cell_index].text}">'

                # PDF:
                for word in cell.text.split(' '):
                    textobject = canvas_.beginText(2 * cm, 29.7*cm - (line_counter%lines_per_page*0.5*cm) - 2*cm)

                    txt_word, index = docx_add_word(word, textobject, indexes_offset+len(txt_))
                    txt_ += txt_word

                    line_counter += 1
                    if line_counter % lines_per_page == lines_per_page-1:
                        canvas_.showPage()

                    canvas_.drawText(textobject)

                cells_added.append(cell._tc)  # for merged cells

                txt_ += '</td>' if add_tags else ''
        txt_ += '</tr>\n' if add_tags else ''
    txt_ += '</table>\n' if add_tags else ''
    return txt_


###################################################################################################


def insert_run_at_position(par, pos, txt=''):
    """
    Insert a new run with text {txt} into paragraph {par}
    at given position {pos}.
    Returns the newly created run.
    """
    p = par._p
    new_run = par.add_run(txt)
    p.insert(pos + 1, new_run._r)

    return new_run


def insert_run_before(par, run, txt=''):
    """
    Insert a new run with text {txt} into paragraph before given {run}.
    Returns the newly created run.
    """
    run_2 = par.add_run(txt)
    run._r.addprevious(run_2._r)

    return run_2


def insert_run_after(par, run, txt=''):
    """
    Insert a new run with text {txt} into paragraph after given {run}.
    Returns the newly created run.
    """
    run_2 = par.add_run(txt)
    run._r.addnext(run_2._r)

    return run_2


def copy_run_format(run_src, run_dst):
    """
    Copy formatting from {run_src} to {run_dst}.
    """
    rPr_target = run_dst._r.get_or_add_rPr()
    rPr_target.addnext(deepcopy(run_src._r.get_or_add_rPr()))
    run_dst._r.remove(rPr_target)


def split_run_by(par, run, split_by):
    """
    Split text in {run} from paragraph {par} by positions
    provided by {split_by}, while retaining original {run}
    formatting.
    Returns list of split runs starting with original {run}.
    """
    txt = run.text
    txt_len = len(txt)

    if not all(isinstance(i, int) for i in split_by):
        raise ValueError("Split positions must be integer numbers")

    split_list = [i if i >= 0 else txt_len + i for i in split_by]

    if not all(split_list[j] <= split_list[j + 1] for j in range(len(split_list) - 1)):
        raise ValueError("Split positions must be sorted to make sense")

    if split_list[0] < 0:
        raise ValueError("A split position cannot be less than -<text length>")

    split_list.insert(0, 0)
    split_list.append(None)
    split_txts = [txt[split_list[i]:split_list[i + 1]] for i in range(len(split_list) - 1)]
    run.text = split_txts[0]
    split_txts.pop(0)
    new_runs = [run]
    for next_txt in split_txts:
        new_runs.append(insert_run_after(par, new_runs[-1], next_txt))
        copy_run_format(run, new_runs[-1])

    return new_runs


###################################################################################################


def docx_highlight(document, annotations, shift_indexes_per_para=False):
    def run_highlight(run):
        rPr = run._r.get_or_add_rPr()
        rPr.highlight_val = WD_COLOR_INDEX.YELLOW

    for anno in annotations:
        index_start = anno['s']
        index_end = anno['e']

        txt_ = ''
        in_highlight = False
        for index_para, para in enumerate(document.paragraphs):
            # If it is a .txt file. It should be shifted because of additional "\n"
            if shift_indexes_per_para:
                index_start += index_para
                index_end += index_para

            for index_run, run in enumerate(para.runs):
                text_main_run = run.text

                is_changed = False

                index_to_split = []
                runs_to_highlight_idx = []
                # The condition is to find the beginning of of the highlighting
                if len(txt_) <= index_start <= len(txt_ + run.text):

                    # the condition is for the situation where the text to be highlighted start at beginning of the run
                    if index_start != len(txt_):
                        index_to_split.append(index_start - len(txt_))
                        runs_to_highlight_idx.append(1)
                    else:
                        runs_to_highlight_idx.append(0)

                    is_changed = True
                    in_highlight = True

                if in_highlight and len(txt_) <= index_end <= len(txt_ + run.text):
                    index_to_split.append(index_end - len(txt_))
                    # print('END', index_end, len(txt_))
                    if not is_changed:
                        runs_to_highlight_idx.append(0)
                    is_changed = True
                    in_highlight = False

                if in_highlight and not is_changed:
                    if index_start < len(txt_) < index_end and index_start < len(txt_ + run.text) < index_end:
                        # run between index_start and index_end
                        # print('MIDDLE', index_start, len(txt_), len(txt_+run.text), index_end)
                        index_to_split.append(0)
                        runs_to_highlight_idx.append(0)

                if len(index_to_split) != 0:
                    runs = split_run_by(para, run, index_to_split)

                    # print('index_to_split:', index_to_split,
                    #       'runs_to_highlight_idx', runs_to_highlight_idx,
                    #       'all_runs:', len(runs))

                    runs_to_highlight = list(map(runs.__getitem__, runs_to_highlight_idx))
                    for i in runs_to_highlight:
                        run_highlight(i)
                txt_ += text_main_run
            txt_ += '\n'

    # document.save('dest1_highlighted.docx')
    return document


def create_docx_from_txt(txt_fullpath):
    document = Document()

    with open(txt_fullpath, 'r', encoding='utf-8') as f:
        p = document.add_paragraph(f.read())
    return document


class TxtToPDF:
    FONT_NAME = 'Helvetica'
    FONT_SIZE = 7
    MAX_PAGE_WIDTH = 300
    MAX_LINES_IN_PAGE = 30

    def __init__(self, text, output_path, annotations=[], highlight_txt=False):
        self.text = text.decode()
        self.canvas_ = canvas.Canvas(output_path, pagesize=pagesizes.A4)
        # self.canvas_.setFillColor(Color(100, 0, 0, alpha=0.2))
        self.page_width = pagesizes.A4[0]
        self.page_height = pagesizes.A4[1]

        self.bboxes = []
        self.result = defaultdict(list)
        self.annotations = annotations
        self.highlight_txt = highlight_txt

        self.line_counter = 0
        self.index_offset = 0
        self.page_number = 1

        self.execute()

        # pprint.pprint(self.bboxes)
        # pprint.pprint(self.annotations)
        # pprint.pprint(self.result)

    def preprocess(self):
        soup = bs4.BeautifulSoup(self.text, 'html.parser')
        text_no_space = str(self.text).replace(' ', '')
        for i in soup.find_all():
            if f'</{i.name}>' not in text_no_space:
                self.text = self.text.replace(f'<{i.name}>', i.name)

    def add_line(self, textobject, words):
        # print('adding each line', words)
        if textobject._canvas.bottomup:
            textobject._y -= textobject._leading
        else:
            textobject._y += textobject._leading
        textobject._y0 = textobject._y

        bboxes = []

        for word in words:
            x_start = textobject.getCursor()[0]
            textobject.textOut(word + ' ')
            x_end = textobject.getCursor()[0]

            index_start = self.index_offset
            index_end = index_start + len(word)
            self.index_offset += len(word) + 1
            bboxes.append({
                'Geometry': {
                    'BoundingBox':
                        {
                            'Left': x_start,
                            'Width': x_end - x_start,
                            'Top': textobject._y,
                            'Height': textobject._leading
                        }
                },
                'Text': word,
                'index_start': index_start,
                'index_end': index_end,
            })

        self.bboxes.extend(bboxes)
        return bboxes

    def add_text(self, text):
        # Split the text into multiple lines compatible with the length of a pdf page
        lines_aligned = utils.simpleSplit(text, fontName=self.FONT_NAME, fontSize=self.FONT_SIZE, maxWidth=self.MAX_PAGE_WIDTH)

        for index_line, line in enumerate(lines_aligned):
            self.line_counter += 1
            textobject = self.canvas_.beginText(2 * cm, 29.7 * cm - (self.line_counter % self.MAX_LINES_IN_PAGE * 0.5 * cm) - 2 * cm)

            bboxes_new = self.add_line(textobject, line.split(' '))

            self.add_bboxes(bboxes=bboxes_new)

            self.canvas_.drawText(textobject)
            if self.line_counter % self.MAX_LINES_IN_PAGE == (self.MAX_LINES_IN_PAGE - 1):
                self.canvas_.showPage()
                self.page_number += 1

    def add_bboxes(self, bboxes):
        for annotation in self.annotations:
            annotation_start = int(annotation.get('s'))
            annotation_end = int(annotation.get('e'))
            blocks = []

            for bbox_ in bboxes:
                indexes_key_start = bbox_['index_start']
                indexes_key_end = bbox_['index_end']

                if (indexes_key_start <= annotation_start <= indexes_key_end) or \
                        (indexes_key_start <= annotation_end <= indexes_key_end) or \
                        (
                                # for the words inside the annotation
                                annotation_start <= indexes_key_start <= annotation_end and
                                annotation_start <= indexes_key_end <= annotation_end
                        ):
                    blocks.append(bbox_)

                    if self.highlight_txt:
                        self.canvas_.rect(
                            bbox_['Geometry']['BoundingBox']['Left'],
                            bbox_['Geometry']['BoundingBox']['Top'] + bbox_['Geometry']['BoundingBox']['Height'],  # buttom left
                            bbox_['Geometry']['BoundingBox']['Width'],
                            bbox_['Geometry']['BoundingBox']['Height'],
                            fill=False,
                        )

            if len(blocks) != 0:
                from app_process.rore.utils import calc_union_blocks, get_ayako_output
                # print(annotation)
                self.result[str(self.page_number)].append(
                    {
                        'lines': self.get_ayako_format(calc_union_blocks(blocks)),
                        # 'blocks': blocks,
                        'type': annotation['type'],
                        'text': ' '.join([i['Text'] for i in blocks])
                    }
                )

    def execute(self):
        self.preprocess()
        soup = bs4.BeautifulSoup(self.text, 'html.parser')

        for element in soup:
            if isinstance(element, bs4.element.Tag) and 'table' in element.name:
                self.add_table(element_table=element)
            else:
                self.add_text(element.get_text())

        self.canvas_.save()

    def add_table(self, element_table):

        def iter_table(element_table):
            for row in element_table:
                if isinstance(row, bs4.element.Tag):
                    for cell in row:
                        if isinstance(cell, bs4.element.Tag):
                            yield cell.get_text()

        for cell_text in iter_table(element_table):
            self.add_text(cell_text)

    def get_result(self):
        return self.result

    def get_ayako_format(self, blocks):
        for block in blocks:
            block.update({
                    "height": self.page_height,
                    "width": self.page_width,
                })
        return blocks


def test_table(data_str):
    test_text = """
    normal_text_0
    <table>
    <tr><td>(0,0)</td><td>(0,1)</td></tr>
    <tr><td>(1,0)</td><td>(1,1)</td></tr>
    </table>
    normal_text_1
    <key>dffdf</key>
    normal_text_2
    <salam>
    <salam2>
    normal_text_3
    <table>
    <tr><td><salam2></td><td>1(0,1)</td></tr>
    <tr><td>1(1,0)</td><td>1(1,1)</td></tr>
    </table>
    normal_text_4
    """

    import bs4
    from bs4 import BeautifulSoup
    # soup = BeautifulSoup(test_text, 'html.parser')
    soup = BeautifulSoup(test_text, 'lxml')

    # for ii in soup.find_all():
    #     print('NEW', ii.name, ii.text)

    soup = BeautifulSoup(test_text, 'html.parser')

    def preprocess(text, soup):
        for i in soup.find_all():
            if f'</{i.name}>' not in text.replace(' ', ''):
                text = text.replace(f'<{i.name}>', i.name)
                print(i.name)
        return text

    text_preproccesed = preprocess(test_text, soup)
    soup = BeautifulSoup(text_preproccesed, 'html.parser')
    # print(soup.get_text())

    for element in soup:
        if isinstance(element, bs4.element.Tag):
            if 'table' in element.name:
                print('table', element)
                for row in element:
                    if isinstance(row, bs4.element.Tag):
                        print('ROW', type(row))
                        for ce in row:
                            if isinstance(ce, bs4.element.Tag):

                                print('CE', type(ce), ce.get_text())

    for i in soup.find_all():
        print(type(i), i)
        if isinstance(i, bs4.element.Tag):
            print(i.name)
